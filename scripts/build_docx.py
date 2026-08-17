# -*- coding: utf-8 -*-
"""Phase 0-3. v2.4 마스터 CSV → 부록 docx (python-docx).

v2.3 부록 형식을 따른다: A4 가로, 재해유형별 그룹 헤더, 8열 표,
Appendix A~D 구조. status=excluded 행은 제외한다.
"""
import csv
import io
import os
import sys
from collections import Counter, OrderedDict

sys.path.insert(0, "scripts")
import ptd_common as C

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

ACC_LABEL = OrderedDict([
    ("떨어짐", "떨어짐 (Falls)"),
    ("무너짐", "무너짐 (Collapse)"),
    ("물체에맞음", "물체에 맞음 (Struck-by falling object)"),
    ("넘어짐", "넘어짐 (Slips & Trips)"),
    ("끼임", "끼임 (Caught-in/between)"),
    ("부딪힘", "부딪힘 (Struck-against)"),
])
TRADE_LABEL = {"거푸집설치": "거푸집 설치", "거푸집해체": "거푸집 해체",
               "타설": "타설", "철근": "철근", "자재운반": "자재 운반"}
HOC_RANK = dict((h, i) for i, h in enumerate(C.HOC_LEVELS))

A_COLS = ["ID", "공종 (Trade)", "PtD 대안 (Directive)", "HoC", "설계결정",
          "Action By", "정량 사양 (Spec)", "출처 (Source)"]
A_WIDTHS = [3.4, 2.2, 10.5, 1.8, 1.4, 2.0, 3.6, 2.6]   # cm, 합 27.5


def landscape_a4(doc):
    s = doc.sections[0]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = Cm(29.7), Cm(21.0)
    for m in ("left_margin", "right_margin"):
        setattr(s, m, Cm(1.1))
    s.top_margin = s.bottom_margin = Cm(1.2)


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def new_table(doc, headers, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        shade(c, "D9E2F3")
        if widths:
            c.width = Cm(widths[i])
    return t


def add_row(t, values, widths=None, size=7.5, bold_first=False):
    cells = t.add_row().cells
    for i, v in enumerate(values):
        cells[i].text = ""
        run = cells[i].paragraphs[0].add_run(v if v else "–")
        run.font.size = Pt(size)
        if bold_first and i == 0:
            run.bold = True
        if widths:
            cells[i].width = Cm(widths[i])
    return cells


def group_row(t, text):
    cells = t.add_row().cells
    m = cells[0]
    for c in cells[1:]:
        m = m.merge(c)
    m.text = ""
    run = m.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    shade(m, "EDEDED")


def para(doc, text, size=8, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    return p


def load_ttl_bits():
    """Appendix B(전거)·D(권고)를 v2.4 TTL 에서 읽는다.

    이전에는 여기서 C.SRC_TTL(v2.3)을 읽어, Appendix A 는 v2.4 CSV 인데
    B·D 만 v2.3 을 보는 교차 버전 상태였다. 지금은 내용이 같아 드러나지
    않지만 버전이 갈라지면 조용히 어긋난다.

    OUT_TTL 이 없으면 SRC_TTL 로 대체하지 않고 중단한다 — 조용한 대체가
    바로 이 결함을 만든 원인이다.
    """
    import rdflib
    from rdflib import RDFS
    if not os.path.exists(C.OUT_TTL):
        raise SystemExit(
            "[중단] %s 가 없습니다. Appendix B·D 는 v2.4 TTL 에서 읽어야 하며\n"
            "        v2.3 으로 조용히 대체하지 않습니다.\n"
            "        먼저 실행하세요: python scripts/build_ttl.py" % C.OUT_TTL)
    g = rdflib.Graph()
    g.parse(C.OUT_TTL, format="turtle")
    P = rdflib.Namespace(C.PTD_NS)
    refs = sorted((C.uri_frag(r), str(g.value(r, P.sourceDocument) or ""))
                  for r in g.subjects(rdflib.RDF.type, P.Reference))
    conf = sorted((C.uri_frag(s), str(g.value(s, RDFS.comment) or ""))
                  for s in g.subjects(rdflib.RDF.type, P.ConflictResolution))
    return refs, conf


def load_archetypes():
    if not os.path.exists(C.KALIS_UNADOPTED):
        return []
    with io.open(C.KALIS_UNADOPTED, encoding=C.OUTPUT_ENCODING, newline="") as f:
        return list(csv.DictReader(f))


def build():
    C.ensure_cwd()
    C.ensure_build()
    allrows = C.read_master()
    rows = [r for r in allrows if r["status"] != "excluded"]
    excl = [r for r in allrows if r["status"] == "excluded"]

    doc = Document()
    landscape_a4(doc)
    doc.styles["Normal"].font.size = Pt(9)

    # ── Appendix A
    doc.add_heading("Appendix A. PtD Alternative Library for RC Frame "
                    "Construction (재해유형별 PtD 대안 라이브러리)", level=1)
    n_t = sum(1 for r in rows if r["promoted"] == "TRUE")
    n_u = sum(1 for r in rows if r["promoted"] == "UNSURE")
    n_f = sum(1 for r in rows if r["promoted"] == "FALSE")
    para(doc, "구성: KnowledgeEntry %d건 (v2.4, 찔림 계열 %d건 범위 제외). "
              "재해유형(사망 점유율 순) 그룹화. "
              "* = 승격 TRUE (%d건), ? = 판정 유보 UNSURE (%d건), "
              "표기 없음 = FALSE (%d건)."
         % (len(rows), len(excl), n_t, n_u, n_f))
    para(doc, "승격 판정은 결정 트리로 부여한 1차 초안이며 확정이 아니다. "
              "UNSURE 는 사람이 판단해야 할 항목의 표시로, 양쪽 논거는 "
              "마스터 CSV 의 adjudication_note 열에 있다.", italic=True)
    para(doc, "음영 행(KE_K_*)은 국토안전관리원 위험요소 프로파일"
              "(2025.09.23 개방본, 전체 55,546건) 스크리닝으로 추가된 항목이다.",
         italic=True)

    t = new_table(doc, A_COLS, A_WIDTHS)
    for acc, label in ACC_LABEL.items():
        sub = [r for r in rows if r["accident_type"] == acc]
        if not sub:
            continue
        kn = sum(1 for r in sub if r["entry_id"].startswith("KE_K_"))
        pn = sum(1 for r in sub if r["promoted"] == "TRUE")
        group_row(t, "%s   총 %d건 / KALIS 신규 %d건 / 승격 %d건"
                  % (label, len(sub), kn, pn))
        sub.sort(key=lambda r: (HOC_RANK.get(r["hoc_level"], 99), r["entry_id"]))
        for r in sub:
            mark = {"TRUE": "*", "UNSURE": "?"}.get(r["promoted"], "")
            cells = add_row(t, [
                r["entry_id"] + mark,
                TRADE_LABEL.get(r["trade"], r["trade"]),
                r["directive_ko"],
                r["hoc_level"],
                "O" if r["design_decidable"] == "TRUE" else "X",
                r["action_by"].replace(";", ", "),
                r["spec"],
                r["source_id"].replace(";", ", "),
            ], A_WIDTHS)
            if r["entry_id"].startswith("KE_K_"):
                for c in cells:
                    shade(c, "FFF2CC")

    orphan = [r for r in rows if r["accident_type"] not in ACC_LABEL]
    if orphan:
        group_row(t, "재해유형 미지정 (%d건)" % len(orphan))
        for r in orphan:
            add_row(t, [r["entry_id"], TRADE_LABEL.get(r["trade"], r["trade"]),
                        r["directive_ko"], r["hoc_level"],
                        "O" if r["design_decidable"] == "TRUE" else "X",
                        r["action_by"].replace(";", ", "), r["spec"],
                        r["source_id"].replace(";", ", ")], A_WIDTHS)

    if excl:
        para(doc, "v2.4 범위 제외: %s (찔림 계열 — status=excluded. 행은 마스터 "
                  "CSV 에 보존되며 TTL·부록에서만 제외된다)"
             % ", ".join(r["entry_id"] for r in excl), italic=True)

    exc = [r for r in rows if r["hoc_rule_exception"] == "TRUE"]
    if exc:
        para(doc, "HoC ↔ rule_type 대응 예외 %d건: %s. 이 대응은 검증 대상 "
                  "가설이지 제약이 아니므로 분류를 수정하지 않고 기록만 한다."
             % (len(exc), ", ".join(r["entry_id"] for r in exc)), italic=True)

    # ── Appendix B
    doc.add_page_break()
    doc.add_heading("Appendix B. Reference List (전거 목록)", level=1)
    refs, conf = load_ttl_bits()
    tb = new_table(doc, ["Ref. ID", "Full citation"], [4.0, 23.5])
    for rid, cit in refs:
        add_row(tb, [rid, cit], [4.0, 23.5], bold_first=True)

    # ── Appendix C
    doc.add_page_break()
    doc.add_heading("Appendix C. 국내 설계단계 저감대책의 통제단계위계 분포",
                    level=1)
    para(doc, "국토안전관리원 위험요소 프로파일의 설계단계 기술을 통제 원형으로 "
              "묶고 HoC 등급을 부여한 결과다. 대상 범위: 시설물분류(대)=건축, "
              "공종분류(중)=철근콘크리트공사·가설공사, 인적피해가 v2.4 대상 "
              "6개 재해유형인 행, 중복 제거 후 3,227건.")

    arch = load_archetypes()
    tot = sum(int(a["count"]) for a in arch) or 1
    byhoc = Counter()
    for a in arch:
        byhoc[a["hoc_level"] or "(미분류)"] += int(a["count"])

    W5 = [7.0, 3.0, 3.0, 4.5, 10.0]
    tc = new_table(doc, ["HoC 등급", "건수", "비율", "누적(상위→하위)", "비고"], W5)
    cum = 0
    order = [h for h in C.HOC_LEVELS if byhoc.get(h)]
    if byhoc.get("(미분류)"):
        order.append("(미분류)")
    for h in order:
        n = byhoc[h]
        cum += n
        add_row(tc, [h, "{:,}".format(n), "%.1f%%" % (100.0 * n / tot),
                     "%.1f%%" % (100.0 * cum / tot),
                     "미채택 원형 집계" if h != "(미분류)" else "원형 미분류·공란"],
                W5)
    add_row(tc, ["합계", "{:,}".format(tot), "100.0%", "100.0%", ""], W5,
            bold_first=True)

    upper = sum(byhoc.get(h, 0) for h in ("위험회피", "제거", "대체"))
    admin = byhoc.get("관리적", 0)
    para(doc, "요약 — 상위 3단(위험회피·제거·대체)의 합계는 %.1f%% 인 반면 "
              "관리적 통제가 %.1f%% 를 차지한다."
         % (100.0 * upper / tot, 100.0 * admin / tot))
    para(doc, "해석 — 설계안전성검토 제도의 표준 참조자료가 위험을 제거·대체하는 "
              "설계 결정보다 위험을 확인하고 문서화하는 절차에 편중되어 있음을 "
              "시사한다. 최대 원형인 구조검토 계열은 작업자의 위치·체류시간을 "
              "바꾸지 않아 실행층으로 승격되지 않는다(NO_EXPOSURE).")

    doc.add_heading("C-1. 통제 원형별 분류 및 라이브러리 반영 결과", level=2)
    W5b = [8.5, 3.0, 3.0, 4.0, 9.0]
    tc1 = new_table(doc, ["통제 원형 (control archetype)", "HoC", "건수",
                          "판정", "라이브러리 반영"], W5b)
    for a in arch:
        add_row(tc1, [a["archetype_label"], a["hoc_level"],
                      "{:,}".format(int(a["count"])),
                      a["reason_code"] or a["promoted"],
                      "미채택" if a["archetype"].startswith("U_") else "채택"],
                W5b)

    # ── Appendix D
    doc.add_page_break()
    doc.add_heading("Appendix D. 정리 권고 (ConflictResolution)", level=1)
    W2 = [3.5, 24.0]
    td = new_table(doc, ["ID", "권고 내용"], W2)
    for cid, txt in conf:
        add_row(td, [cid, txt], W2, bold_first=True)

    doc.save(C.OUT_DOCX)
    print("docx 생성 : %s" % C.OUT_DOCX)
    print("  용지            : A4 가로 (29.7 x 21.0 cm)")
    print("  Appendix A 항목 : %d (제외 %d)" % (len(rows), len(excl)))
    print("  Appendix B 전거 : %d" % len(refs))
    print("  Appendix C 원형 : %d (대상 %s행)" % (len(arch), "{:,}".format(tot)))
    print("  Appendix D 권고 : %d" % len(conf))
    return 0


if __name__ == "__main__":
    sys.exit(build())
